# -*- coding: utf-8 -*-
"""将 Markdown 报告转换为带公式(图片)、表格、图片的 Word 文档。"""
import re, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SRC = "/workspace/风险定价模型报告.md"
FIG = "/workspace/figs"
OUT = "/workspace/风险定价模型报告.docx"

doc = Document()
# 基础字体
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(10.5)

def add_para(text, size=10.5, bold=False, align=None, color=None, italic=False):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = RGBColor(*color)
    return p

def render_inline(text):
    """处理行内 **加粗** 与 `代码`，LaTeX $...$ 保留为普通文本（Word 公式另以图片呈现）。"""
    # 先拆出代码段与加粗
    parts = re.split(r'(\*\*.*?\*\*|`[^`]*`)', text)
    runs = []
    for seg in parts:
        if seg.startswith("**") and seg.endswith("**"):
            runs.append((seg[2:-2], True, False))
        elif seg.startswith("`") and seg.endswith("`"):
            runs.append((seg[1:-1], False, True))
        else:
            runs.append((seg, False, False))
    return runs

def add_rich(text, size=10.5, align=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    for txt, b, code in render_inline(text):
        r = p.add_run(txt)
        r.font.size = Pt(size); r.bold = b
        if code: r.font.name = "Consolas"
    return p

# 公式图片：用 matplotlib 渲染 LaTeX（若可用），否则保留原文
FORMULA_DIR = "/workspace/figs/_formulas"
os.makedirs(FORMULA_DIR, exist_ok=True)
try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    _has_mpl = True
except Exception:
    _has_mpl = False

_formula_cache = {}
def formula_to_png(tex, key):
    if not _has_mpl: return None
    if key in _formula_cache: return _formula_cache[key]
    try:
        fig = plt.figure(figsize=(7.2, 0.55))
        fig.text(0.01, 0.5, f"${tex}$", fontsize=12, va="center")
        path = os.path.join(FORMULA_DIR, f"f{key}.png")
        fig.savefig(path, dpi=160, bbox_inches="tight", transparent=True)
        plt.close(fig)
        _formula_cache[key] = path
        return path
    except Exception:
        return None

with open(SRC, encoding="utf-8") as f:
    lines = f.read().splitlines()

i = 0
fkey = 0
img_re = re.compile(r'^!\[(.*?)\]\((.*?)\)\s*$')
hdr_re = re.compile(r'^(#{1,4})\s+(.*)$')
tbl_sep = re.compile(r'^\|[\s\-:|]+\|$')
in_code = False

while i < len(lines):
    line = lines[i]
    if line.strip().startswith("```"):
        in_code = not in_code
        i += 1; continue
    if in_code:
        add_para(line, size=9, color=(60,60,60)); i += 1; continue
    m = img_re.match(line)
    if m:
        cap, path = m.group(1), m.group(2)
        fp = os.path.join("/workspace", path) if not os.path.isabs(path) else path
        if os.path.exists(fp):
            doc.add_picture(fp, width=Inches(5.4))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if cap:
                c = add_para(cap, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(110,110,110))
        i += 1; continue
    m = hdr_re.match(line)
    if m:
        level = len(m.group(1)); txt = m.group(2).strip()
        p = doc.add_heading(level=level)
        r = p.add_run(txt); r.font.name = "Microsoft YaHei"
        i += 1; continue
    if line.strip().startswith("> "):
        add_para(line.strip()[2:], size=10, italic=True, color=(90,90,90)); i += 1; continue
    if tbl_sep.match(line):
        # 收集表格
        rows = []
        j = i
        while j < len(lines) and lines[j].strip().startswith("|"):
            cells = [c.strip() for c in lines[j].strip().strip("|").split("|")]
            rows.append(cells); j += 1
        if len(rows) >= 2:
            header = rows[0]; body = rows[2:]  # 跳过分隔行
            t = doc.add_table(rows=1, cols=len(header))
            t.style = "Light Grid Accent 1"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for c, h in enumerate(header):
                cell = t.rows[0].cells[c]
                cell.text = ""
                run = cell.paragraphs[0].add_run(h); run.bold = True; run.font.size = Pt(9)
            for br in body:
                cells = t.add_row().cells
                for c, val in enumerate(br):
                    cells[c].text = ""
                    run = cells[c].paragraphs[0].add_run(val); run.font.size = Pt(9)
        i = j; continue
    if line.strip() == "":
        i += 1; continue
    # 行内公式 $...$ 单独成段则渲染为图片，否则富文本
    if line.strip().startswith("$") and line.strip().endswith("$") and line.count("$")==2:
        tex = line.strip().strip("$")
        fkey += 1
        png = formula_to_png(tex, fkey)
        if png:
            doc.add_picture(png, width=Inches(5.6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            add_rich(line.strip(), size=10.5)
        i += 1; continue
    add_rich(line, size=10.5); i += 1

doc.save(OUT)
print("已生成:", OUT, "段落数:", len(doc.paragraphs))
